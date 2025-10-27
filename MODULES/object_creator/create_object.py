
from openpyxl import load_workbook

from docx import Document
from docx2pdf import convert
from ENV.environment import pegar_modelos, atualizar_numero_protocolo
from MODULES.document_dealer.doc_dealer import DocHelper

import locale
import os
import shutil

class Termo:

    index = 1 

    def __init__(self, contrato, contratado, objeto, af, mensagem, gestor):
        self.contratado = contratado 
        self.contrato = contrato 
        self.mensagem = mensagem
        self.objeto = objeto 
        self.af = af
        self.gestor = gestor 

    def setRelatorioInfo(self, liquidacao, valor, data):
        self.liquidacao = liquidacao 
        self.valor = valor 
        self.data = data 

    def copiar_arquivo(self, arquivo, numero_protocolo = None, mesmo_protocolo = False):

        if (arquivo == 'termo'): 
            nome_arquivo = f'{Termo.index}. {self.contratado} - AF {self.af[:-4]}'
            endereco_copia = os.getcwd() + rf"\{nome_arquivo}" + ".docx"
            Termo.index += 1 
        elif (arquivo == 'protocolo'): 
            endereco_copia = os.getcwd() + rf"\Protocolo N° {numero_protocolo} - Tesouraria.docx"
        
        endereco_modelo = pegar_modelos(arquivo)
       
        if (mesmo_protocolo == False): shutil.copy(endereco_modelo, endereco_copia)

        return endereco_copia
    
    def criar_termo(self, tipo):
        
        termo = self.copiar_arquivo("termo")

        doc = Document(termo)

        def definir_tabela():

            tabela = doc.tables[0]
             
            if "credenciamento" in self.objeto.lower():
                campo = "CREDENCIAMENTO N°"
            elif tipo["corresponde_a"] == "ata":
                campo = "ATA N°"
            else:
                campo = "CONTRATO N°" 


            dados = [{"celula": (1, 0), "conteudo": campo, "negrito": True}, 
                     {"celula": (1, 1), "conteudo": self.contrato, "negrito": False }, 
                     {"celula": (2, 1), "conteudo": self.contratado, "negrito": False }, 
                     {"celula": (3, 1), "conteudo": self.objeto, "negrito": False }, 
                     {"celula": (4, 1), "conteudo": self.af, "negrito": False}, 
                     {"celula": (6, 1), "conteudo": self.mensagem, "negrito": False }, 
                    ]

            DocHelper.modificar_tabela(tabela, dados)

        def adicionar_espaco():
            doc.add_paragraph("")
        
        def adicionar_data():
            data = DocHelper.encontrar_data_de_hoje_em_extenso()
            paragrafo = doc.add_paragraph()
            DocHelper.criar_texto(paragrafo, data, negrito=True, posicionamento = "Direita", fonte = "Cambria")
            pass 

        def adicionar_assinatura():
            paragrafo = doc.add_paragraph()
            DocHelper.adicionar_linha_de_assinatura(paragrafo, self.gestor)

        definir_tabela()
        adicionar_espaco()
        adicionar_data()
        adicionar_espaco()
        adicionar_assinatura()

        doc.save(termo)
        Termo.salvar_pdf(termo[:-5])
        
    def salvar_pdf(docx):
        pdf = docx.replace("WORD", "PDF") + ".pdf"
        docx = docx + ".docx"
        convert(docx, pdf)

    @staticmethod
    def criar_relatorio(termos, numero_protocolo, mesmo_protocolo):
        
        protocolo = Termo.copiar_arquivo(None, "protocolo", numero_protocolo, mesmo_protocolo)
        doc = Document(protocolo)

        def converter_currency(valor):
            locale.setlocale(locale.LC_ALL, "pt_BR.UTF-8")
            return locale.currency(float(valor), grouping =True)
        
        def adicionar_protocolo():
            substituir_protocolo = doc.paragraphs[1]
            substituir_protocolo.text = ""
            DocHelper.criar_texto(substituir_protocolo, f"PROTOCOLO DE RECEBIMENTO - NÚMERO {numero_protocolo}", negrito = True)
   
        def criar_tabela():
            
            tabela = doc.tables[0]
            for i in range(len(termos)):
                
                nova_linha = tabela.add_row()
            
                coluna_um = nova_linha.cells[0].paragraphs[0]
                coluna_dois = nova_linha.cells[1].paragraphs[0]
                coluna_tres = nova_linha.cells[2].paragraphs[0]
                coluna_quatro = nova_linha.cells[3].paragraphs[0]
                
            
                DocHelper.criar_texto(coluna_um, termos[i].contratado,  px = 8, negrito = True, fonte = "Arial")
                DocHelper.criar_texto(coluna_dois, termos[i].liquidacao,  px = 8, negrito = True, fonte = "Arial")
                DocHelper.criar_texto(coluna_tres, termos[i].data,  px = 8, negrito = True, fonte = "Arial")
                DocHelper.criar_texto(coluna_quatro, converter_currency(termos[i].valor),  px = 8, negrito = True, fonte = "Arial")
                
           
        adicionar_protocolo() 
        criar_tabela()

        ask_it = str(input("Atualizar o protocolo? (Y/N) --> ")).lower()

        if ask_it == "y": atualizar_numero_protocolo()
        doc.save(protocolo)

        pass 

def capturar_info_planilha(localizacao_planilha):

    def formatar_data(objeto):
        return objeto.date().strftime("%d/%m/%Y")
    
    def customizar_mensagem(tipo_nota):
        if (tipo_nota == "Locação"): return "Por este instrumento, em caráter DEFINITIVO, atestamos que a locação acima identificada atende às exigências contratuais."
        return f"Por este instrumento, em caráter DEFINITIVO, atestamos que os {tipo_nota.lower()} acima identificados atendem às exigências contratuais."
    
    PLANILHA = load_workbook(localizacao_planilha, data_only = True)
    TERMOS = []

    for nome in PLANILHA.sheetnames: 
        if nome != "Base" and nome != "Fiscais":

            SHEET = PLANILHA[nome]

            if (SHEET["A4"].value == None): 
                continue 
             
            contrato = SHEET["E4"].value 
            contratado = SHEET["B4"].value
            mensagem = customizar_mensagem(str(SHEET["D16"].value))
            objeto = SHEET["F4"].value 
            af = SHEET["A20"].value

            liquidacao = SHEET["A8"].value 
            data_liquidacao = formatar_data(SHEET["B12"].value)
            valor = SHEET["C12"].value

            if ("ATA" in localizacao_planilha):
                gestor = "MURILO SOARES DE OLIVEIRA\nGESTOR DE ATA"
            else: 
                gestor = "RONALDO DE SOUZA MARCILIO\nGESTOR DE CONTRATO"
        
            novo_termo = Termo(contrato, contratado, objeto, af, mensagem, gestor)
            novo_termo.setRelatorioInfo(liquidacao, valor, data_liquidacao)

            TERMOS.append(novo_termo)

    return TERMOS